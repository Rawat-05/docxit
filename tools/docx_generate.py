import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)


def add_kv(doc, label, value):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value or "")


def add_device_table(doc, devices):
    if not devices:
        doc.add_paragraph("No devices added.")
        return
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Device Name"
    hdr[2].text = "Type"
    for i, d in enumerate(devices, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = d.get("name", "")
        row[2].text = d.get("type", "")


def generate_doc(data, out_path):
    doc = Document()
    add_heading(doc, "DOCXIT - Documentation Suite", level=1)
    add_heading(doc, "Case Summary", level=2)
    add_kv(doc, "Case Information", data.get("caseInformation", ""))
    add_kv(doc, "Date of Search", data.get("dateOfSearch", ""))
    add_kv(doc, "Conclusion Date", data.get("conclusionDate", ""))
    add_kv(doc, "Party Name", data.get("partyName", ""))
    add_kv(doc, "Authorized Officer", data.get("authorizedOfficer", ""))
    add_kv(doc, "Examiner", data.get("examinerName", ""))

    add_heading(doc, "Devices", level=2)
    add_device_table(doc, data.get("deviceList", []))

    doc.add_page_break()
    add_heading(doc, "Custom Document", level=2)
    custom = data.get("customTemplate", "")
    doc.add_paragraph(custom or "(No custom template provided)")

    doc.save(out_path)


def main():
    if len(sys.argv) < 3:
        print(json.dumps({"error": "Usage: docx_generate.py <data.json> <output.docx>"}))
        return 1
    data_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])
    data = json.loads(data_path.read_text(encoding="utf-8"))
    generate_doc(data, out_path)
    print(json.dumps({"ok": True, "output": str(out_path)}))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
